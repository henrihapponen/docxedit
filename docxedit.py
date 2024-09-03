import docx


def show_line(doc: object, current_text: str, show_errors: bool = True):
    """
    Shows the 'line' of text in the doc where the string is found (without replacing anything).
    Args:
        doc (Object): The docx document object.
        current_text (str): The string to search for.
        show_errors (bool): Whether to show errors or not. Default is True.
    Returns:
        The line of text where the string is found, if it is found.
    """

    try:
        for paragraph in doc.paragraphs:
            if current_text in paragraph.text:
                inline = paragraph.runs

                for i in range(len(inline)):
                    if current_text in inline[i].text:
                        text = inline[i].text
                        inline[i].text = text
                print(f'Text found in line: {paragraph.text}')
            else:
                if show_errors:
                    print(f'Error: {current_text} not found in document.')
    except Exception as e:
        print(f'Error: An Exception occurred: {e}')


def replace_string(doc: object, old_string: str, new_string: str,
                   include_tables: bool = True, show_errors: bool = False):
    """
    Replaces an old string (placeholder) with a new string
    without changing the formatting of the text.
    Args:
        doc (Object): The docx document object.
        old_string (str): The old string to replace.
        new_string (str): The new string to replace the old one with.
        include_tables (bool): Whether to include tables or not. Default is True.
        show_errors (bool): Whether to show errors or not. Default is False.
    Returns:
        Success or Error.
    """

    string_instances_replaced = 0

    for paragraph in doc.paragraphs:
        if old_string in paragraph.text:
            inline = paragraph.runs

            for i in range(len(inline)):
                if old_string in inline[i].text:
                    text = inline[i].text.replace(str(old_string), str(new_string))
                    inline[i].text = text
                    string_instances_replaced += 1
            print(f'Success: Replaced the string "{old_string}" with "{new_string}" in a paragraph')
        else:
            if show_errors:
                print(f'Error: Could not find the string "{old_string}" in the document')

    if include_tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_string in paragraph.text:
                            inline = paragraph.runs

                            for i in range(len(inline)):
                                if old_string in inline[i].text:
                                    text = inline[i].text.replace(str(old_string), str(new_string))
                                    inline[i].text = text
                            string_instances_replaced += 1
                            print(f'Success: Replaced the string "{old_string}" '
                                  f'with "{new_string}" in a table')
                        else:
                            if show_errors:
                                print(f'Error: Could not find {old_string} in a table')

    print(f'Summary: Replaced {string_instances_replaced} instances of "{old_string}" '
          f'with "{new_string}"')


def replace_string_up_to_paragraph(doc: object, old_string: str, new_string: str,
                                   paragraph_number: int, show_errors: bool = True):
    """
    Replaces an old string (placeholder) with a new string
    without changing the format of the text but only up to
    a specific paragraph number.
    Args:
        doc (Object): The docx document object.
        old_string (str): The old string to replace.
        new_string (str): The new string to replace the old one with.
        paragraph_number (int): The paragraph number to stop at.
        show_errors (bool): Whether to show errors or not. Default is True.
    """

    string_instances_replaced = 0

    for index, paragraph in enumerate(doc.paragraphs):
        if index < paragraph_number:
            if old_string in paragraph.text:
                inline = paragraph.runs

                for i in range(len(inline)):
                    if old_string in inline[i].text:
                        text = inline[i].text.replace(str(old_string), str(new_string))
                        inline[i].text = text

                string_instances_replaced += 1

                print(
                    f'Success: Replaced the string "{old_string}" with "{new_string}" '
                    f'in paragraph number {index}')

    print(f'Summary: Replaced {string_instances_replaced} instances of "{old_string}" '
          f'with "{new_string}"')


def remove_paragraph(paragraph: object, show_errors: bool = True):
    """
    Remove a paragraph. Input must be a paragraph object.
    Args:
        paragraph (Object): The paragraph object to remove.
        show_errors (bool): Whether to show errors or not. Default is True.
    Returns:
        Success or Error.
    """

    try:
        paragraph_element = paragraph._element
        paragraph_element.getparent().remove(paragraph_element)
        paragraph._p = paragraph._element = None
        return print(f'Success: Removed paragraph {paragraph}')
    except Exception as e:
        if show_errors:
            return print(f'Error: Could not remove paragraph {paragraph}: {e}')


def remove_lines(doc, first_line: str, number_of_lines: int, show_errors: bool = True):
    """
    Remove a line including any keyword (first_line),
    and a certain number of rows after that.
    Args:
        doc (Object): The docx document object.
        first_line (String): The first line to remove.
        number_of_lines (Integer): The number of lines to remove.
        show_errors (bool): Whether to show errors or not. Default is True.
    Returns:
        Success or Error.
    """

    list_of_paragraphs = []

    for i in doc.paragraphs:
        list_of_paragraphs.append(i)

    for index, i in enumerate(doc.paragraphs):
        if first_line in i.text:
            try:
                remove_paragraph(i)
            except AttributeError:
                if show_errors:
                    print(f'Error: Could not remove line {index}: {i.text}')

            b_var = 0
            c_var = 0
            while b_var < number_of_lines:
                try:
                    print(list_of_paragraphs[index + 1 + b_var])
                    remove_paragraph(list_of_paragraphs[index + 1 + c_var])
                    b_var += 1
                    print(f'Success: Removed line {str(index + 1 + c_var)}')
                except Exception as e:
                    if show_errors:
                        print(f'Error: Could not remove line {str(index + 1 + c_var)} '
                              f'due to exception: {e}')
                    c_var += 1
                    continue


def add_text_in_table(table, row_num: int, column_num: int,
                      new_string: str, show_errors: bool = True):
    """
    Add text to a cell in a table.
    Args:
        table (Object): The table object.
        row_num (int): The row number to add the text to.
        column_num (int): The column number to add the text to.
        new_string (str): The text to add.
        show_errors (bool): Whether to show errors or not. Default is False.
    Returns:
        Success or Error.
    """

    try:
        table.cell(row_num, column_num).text = new_string
        return print(f'Success: Added {new_string} to row {row_num} and column {column_num}')
    except Exception as e:
        if show_errors:
            return print(f'Error: Could not add {new_string} to row {row_num} '
                         f'and column {column_num} due to exception: {e}')


def change_table_font_size(table: object, font_size: int, show_errors: bool = True):
    """
    Change the font size of the whole table.
    Args:
        table (Object): The table object.
        font_size (Integer): The font size to change to.
        show_errors (bool): Whether to show errors or not. Default is True.
    Returns:
        Success or Error.
    """

    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = docx.shared.Pt(font_size)
        return print(f'Success: Changed font size to {font_size}')
    except Exception as e:
        if show_errors:
            return print(f'Error: Could not change font size to {font_size}: {e}')
