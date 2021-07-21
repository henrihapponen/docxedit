# Useful functions to edit Word Documents with Python-Docx module

# Import
from docx import Document
from docx.shared import Pt


# Function definitions

def show_line(current_text):
    """Shows the 'line' of text in the doc where the string is found (without replacing anything)."""

    global doc

    for p in doc.paragraphs:
        if current_text in p.text:
            inline = p.runs

            for i in range(len(inline)):
                if current_text in inline[i].text:
                    text = inline[i].text
                    inline[i].text = text
            print(p.text)

    return


def replace_string(old_text, new_text):
    """Replaces an old string (placeholder) with a new string without changing the formatting of the text."""

    global doc

    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs

            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(str(old_text), str(new_text))
                    inline[i].text = text

    return


def replace_string_up_to_paragraph(old_text, new_text, paragraph_number):
    """
    Replaces an old string (placeholder) with a new string without changing the format of the text
    but only up to a specific paragraph number.
    """

    global doc

    for index, p in enumerate(doc.paragraphs):
        
        # Replace every instance before paragraph number 'paragraph_number'
        if index < paragraph_number:

            if old_text in p.text:
                inline = p.runs

                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(str(old_text), str(new_text))
                        inline[i].text = text

    return


def delete_paragraph(paragraph):
    """Delete a paragraph. Input must be a paragraph object."""

    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

    return


def remove_lines(first_line, number_of_lines):
    """Remove a line including any keyword (first_line), and a certain number of rows after that."""

    list_of_paragraphs = []

    for i in doc.paragraphs:
        list_of_paragraphs.append(i)

    for index, i in enumerate(doc.paragraphs):
        if first_line in i.text:
            try:
                delete_paragraph(i)
            except AttributeError:
                print('Could not remove line ' + str(index) + ': ' + i.text)

            b = 0
            c = 0
            while b < number_of_lines:
                try:
                    print(list_of_paragraphs[a + 1 + b])
                    delete_paragraph(list_of_paragraphs[index + 1 + c])
                    b += 1
                except AttributeError:
                    print('Could not remove line ' + str(index + 1 + c))
                    c += 1
                    continue

    return


def add_text_in_table(table, row_num, column_num, new_text):
    """Add text to a cell in a table."""

    table.cell(row_num, column_num).text = new_text

    return


def change_table_font_size(table, font_size):
    """Change the font size af the whole table."""

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(font_size)

    return


if __name__ == '__main__':

    # For example:
    
    document_path = 'your full document path here'
    doc = Document(document_path)
    
    show_line('Section A')
    replace_string('placeholder', 'new text')
    replace_string_up_to_paragraph('placeholder', 'new text', 10)
    remove_lines('remove this line', 1)
    remove_lines('remove the next 5 lines', 5)
